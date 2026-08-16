from pathlib import Path
import json
import re

from docx import Document
from docx.document import Document as _Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl


MAX_CHARS = 4000
OVERLAP_CHARS = 400

DOCUMENT_PATH = Path("backend/data/documents/36413-j20.docx")
OUTPUT_PATH = Path("backend/data/processed/ts_36_413_chunks.json")


def iter_block_items(document):
    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield Table(child, document)


def table_to_text(table: Table) -> str:
    rows = []

    for row in table.rows:
        cells = []

        for cell in row.cells:
            text = " ".join(
                paragraph.text.strip()
                for paragraph in cell.paragraphs
                if paragraph.text.strip()
            )

            if text:
                cells.append(text)

        if cells:
            rows.append(" | ".join(cells))

    return "\n".join(rows)


def extract_metadata(document: Document) -> dict:
    text = "\n".join(
        paragraph.text.strip()
        for paragraph in document.paragraphs[:20]
        if paragraph.text.strip()
    )

    match = re.search(
        r"3GPP\s+TS\s+(\d+\.\d+)\s+V(\d+\.\d+\.\d+)",
        text,
        re.IGNORECASE,
    )

    if not match:
        raise ValueError("Could not identify specification and version.")

    release_match = re.search(
        r"\(Release\s+(\d+)\)",
        text,
        re.IGNORECASE,
    )

    return {
        "specification": f"TS {match.group(1)}",
        "version": f"V{match.group(2)}",
        "release": release_match.group(1) if release_match else "unknown",
        "source_file": DOCUMENT_PATH.name,
    }


def extract_blocks(document: Document) -> tuple[list[dict], dict]:
    metadata = extract_metadata(document)

    blocks = []
    started = False

    for block in iter_block_items(document):

        if isinstance(block, Paragraph):
            text = block.text.strip()

            if not text:
                continue

            style = block.style.name if block.style else "Normal"

            if text.lower() == "foreword" and style == "Heading 1":
                started = True

            if not started:
                continue

            blocks.append({
                "text": text,
                "style": style,
            })

        elif isinstance(block, Table):

            if not started:
                continue

            text = table_to_text(block)

            if text.strip():
                blocks.append({
                    "text": text,
                    "style": "Table",
                })

    return blocks, metadata


def is_heading(style: str) -> bool:
    return style.startswith("Heading ")


def heading_level(style: str) -> int:
    try:
        return int(style.split()[-1])
    except (ValueError, IndexError):
        return 0


def build_chunks(blocks: list[dict]) -> list[dict]:
    chunks = []
    heading_stack = {}
    current_content = []

    def flush():
        if not current_content:
            return

        text = "\n".join(current_content).strip()

        if not text:
            current_content.clear()
            return

        start = 0

        while start < len(text):
            end = min(start + MAX_CHARS, len(text))
            chunk_text = text[start:end].strip()

            if chunk_text:
                section_path = " > ".join(
                    heading_stack[level]
                    for level in sorted(heading_stack)
                )

                chunks.append({
                    "text": chunk_text,
                    "section_path": section_path,
                })

            if end >= len(text):
                break

            start = end - OVERLAP_CHARS

        current_content.clear()

    for block in blocks:
        text = block["text"]
        style = block["style"]

        if is_heading(style):
            flush()

            level = heading_level(style)

            for existing_level in list(heading_stack):
                if existing_level >= level:
                    del heading_stack[existing_level]

            heading_stack[level] = text

        else:
            current_content.append(text)

    flush()

    return chunks


def process_document() -> list[dict]:
    if not DOCUMENT_PATH.exists():
        raise FileNotFoundError(
            f"Document not found: {DOCUMENT_PATH}"
        )

    document = Document(DOCUMENT_PATH)

    blocks, metadata = extract_blocks(document)
    chunks = build_chunks(blocks)

    for chunk_id, chunk in enumerate(chunks):
        chunk.update({
            "chunk_id": chunk_id,
            "specification": metadata["specification"],
            "version": metadata["version"],
            "release": metadata["release"],
            "source_file": metadata["source_file"],
        })

    return chunks


if __name__ == "__main__":
    chunks = process_document()

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)

    with OUTPUT_PATH.open("w", encoding="utf-8") as file:
        json.dump(
            chunks,
            file,
            ensure_ascii=False,
            indent=2,
        )

    print(f"Total chunks: {len(chunks)}")
    print(f"Saved to: {OUTPUT_PATH}")